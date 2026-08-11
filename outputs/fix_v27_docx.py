# -*- coding: utf-8 -*-
"""Patch V2.6.docx -> V2.7.docx to match the corrected chapter-4 architecture."""
import copy, io
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:\Users\chenhong\WorkBuddy\2026-06-26-14-49-33\szmgzc_pm\数媒智创平台_PRD_V2.6.docx"
DST = r"C:\Users\chenhong\WorkBuddy\2026-06-26-14-49-33\szmgzc_pm\数媒智创平台_PRD_V2.7.docx"

doc = Document(SRC)
body = doc.element.body

def get_text(el):
    return "".join(t.text or "" for t in el.iter(qn('w:t')))

def set_text(p_el, text):
    for r in p_el.findall(qn('w:r')):
        p_el.remove(r)
    r = OxmlElement('w:r'); t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t); p_el.append(r)

def set_style(p_el, style_val):
    pPr = p_el.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr'); p_el.insert(0, pPr)
    ps = pPr.find(qn('w:pStyle'))
    if ps is None:
        ps = OxmlElement('w:pStyle'); pPr.append(ps)
    ps.set(qn('w:val'), style_val)

def find_para(text_startswith):
    for p in body.iter(qn('w:p')):
        if get_text(p).startswith(text_startswith):
            return p
    return None

def insert_para_before(ref_el, text, style=None):
    new_p = OxmlElement('w:p')
    if style:
        pPr = OxmlElement('w:pPr'); ps = OxmlElement('w:pStyle'); ps.set(qn('w:val'), style); pPr.append(ps); new_p.append(pPr)
    r = OxmlElement('w:r'); t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text; r.append(t); new_p.append(r)
    ref_el.addprevious(new_p)
    return new_p

def insert_table_before(ref_el, data):
    tbl = doc.add_table(rows=0, cols=len(data[0]))
    try:
        tbl.style = 'Table Grid'
    except Exception:
        pass
    for row in data:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    ref_el.addprevious(tbl._tbl)
    return tbl

# ---------- 1. version ----------
vp = find_para("文档版本：")
if vp is not None:
    set_text(vp, get_text(vp).replace("文档版本：V2.6", "文档版本：V2.7", 1))

# ---------- 2. delete standalone 4.3 热点榜单 section (incl its table) ----------
start_el = find_para("4.3 热点榜单")
end_el = find_para("4.4 热点助手")
if start_el is None or end_el is None:
    raise SystemExit("4.3/4.4 anchors not found")
# collect children between start (inclusive) and end (exclusive)
to_remove = []
after = False
for child in list(body):
    if child is start_el:
        after = True
        to_remove.append(child); continue
    if child is end_el:
        break
    if after:
        to_remove.append(child)
for el in to_remove:
    body.remove(el)
print("removed 4.3 热点榜单 block elements:", len(to_remove))

# ---------- 2.5 delete redundant 4.4.6 与热点榜单的关系 (merged away) ----------
s6 = find_para("4.4.6 与热点榜单的关系")
e6 = find_para("4.5 生产工具")
if s6 is not None and e6 is not None:
    blk = []; after = False
    for child in list(body):
        if child is s6:
            after = True; blk.append(child); continue
        if child is e6:
            break
        if after:
            blk.append(child)
    for el in blk:
        body.remove(el)
    print("removed 4.4.6 与热点榜单的关系 block:", len(blk))
else:
    print("4.4.6 anchor not found, skip")

# ---------- 3. renumber remaining chapter-4 headers (by exact text) ----------
renumber = [
    ("4.4 热点助手", "4.3 热点助手"),
    ("4.4.1 功能定位", "4.3.1 功能定位"),
    ("4.4.2 页面结构", "4.3.2 页面结构"),
    ("4.4.3 热点搜索", "4.3.3 热点搜索"),
    ("4.4.4 热点预测", "4.3.4 热点预测"),
    ("4.4.5 事件详情", "4.3.6 事件详情"),
    ("4.5 生产工具", "4.4 生产工具"),
    ("4.5.1 智能体", "4.4.1 智能体"),
    ("4.5.2 内容创作", "4.4.2 内容创作"),
    ("4.5.3 智能配音", "4.4.3 智能配音"),
    ("4.5.4 数字主播", "4.4.4 数字主播"),
    ("4.5.5 实用工具", "4.4.5 实用工具"),
    ("4.6 技能市场（SkillHub）", "4.5 技能市场（SkillHub）"),
    ("4.6.1 产品定位", "4.5.1 产品定位"),
    ("4.6.2 核心功能", "4.5.2 核心功能"),
    ("4.6.3 典型技能示例", "4.5.3 典型技能示例"),
    ("4.8 个人中心", "4.6 个人中心"),
]
for p in body.iter(qn('w:p')):
    tx = get_text(p)
    for old, new in renumber:
        if tx == old:
            set_text(p, new); break

# ---------- 4. fold 4.7 历史记录 -> 4.1.5 (move block + rename) ----------
h_start = find_para("4.7 历史记录")
h_end = find_para("4.8 个人中心")  # original text still present (renumber didn't touch 4.7/4.8 except 4.8->4.6 already applied!)
# NOTE: 4.8 already renumbered to 4.6 above, so h_end anchor must use "4.6 个人中心"
h_end = find_para("4.6 个人中心")
if h_start is None or h_end is None:
    raise SystemExit("history anchors not found")
# collect history block
hist = []
after = False
for child in list(body):
    if child is h_start:
        after = True; hist.append(child); continue
    if child is h_end:
        break
    if after:
        hist.append(child)
# detach
for el in hist:
    body.remove(el)
# insert after 4.1.4 第三方工具入口's following sibling before 4.2 知识库
anchor = find_para("4.2 知识库")
for el in hist:
    anchor.addprevious(el)
# rename within moved block
for el in hist:
    if el.tag == qn('w:p'):
        tx = get_text(el)
        if tx == "4.7 历史记录（原\"历史对话\"）":
            set_text(el, "4.1.5 历史记录（原\"历史对话\"）"); set_style(el, "Heading3")
        elif tx == "4.7.1 功能结构":
            set_text(el, "功能结构")
        elif tx == "4.7.2 交互说明":
            set_text(el, "交互说明")
print("folded 历史记录 -> 4.1.5")

# ---------- 5. insert 4.3.5 国内热榜与国际热榜 before 4.3.6 事件详情 ----------
ref = find_para("4.3.6 事件详情")
if ref is None:
    raise SystemExit("4.3.6 anchor not found")
insert_para_before(ref, "4.3.5 国内热榜与国际热榜")
insert_para_before(ref, "国内热榜与国际热榜两个页面即由原独立的「热点榜单」模块合并而来，面向编辑与运营人员提供国内外多平台实时热搜聚合，支撑选题发现与热点跟进。")
insert_para_before(ref, "内容结构")
insert_table_before(ref, [
    ["页面", "覆盖平台", "数据维度"],
    ["国内", "统一覆盖 16 个平台源：百度热搜、腾讯新闻、澎湃新闻、凤凰网、今日头条、抖音、快手、微博、知乎、财联社电报、财联社深度、财联社热榜、哔哩哔哩热搜、哔哩哔哩排行榜、哔哩哔哩热门视频、虎扑", "各平台 Top10 热榜条目，含标题、序号、热度与更新时间；部分平台（如财联社）提供电报/深度/热榜等多形态内容"],
    ["国际", "独立国际热点源集合（与国内平台源不同），共 5 个平台源：参考消息、华尔街见闻快讯、华尔街见闻新闻、华尔街见闻热榜、卫星通讯社（均为可国内访问的国际资讯源）", "国际时事、财经、地缘政治类 Top10"],
])
insert_para_before(ref, "注：国内/国际为两个独立页面，平台源集合不同；国际源并未下架，V2.1 版「国际源已下架」为线上核查误判，V2.2 已订正推断方向，V2.3 据用户提供的线上确切数据将国际源锁定为 5 个平台（参考消息、华尔街见闻快讯/新闻/热榜、卫星通讯社）。")
insert_para_before(ref, "交互说明")
insert_para_before(ref, "在热点助手页切换「国内热榜 / 国际热榜」页面（或访问 /hotspot-assistant?tab=domestic / ?tab=international）")
insert_para_before(ref, "每个榜单卡片显示平台Logo、更新时间、Top10列表")
insert_para_before(ref, "支持刷新按钮更新榜单")
insert_para_before(ref, "支持收藏/更多操作按钮（具体功能需进一步确认）")
print("inserted 4.3.5 国内热榜与国际热榜")

# ---------- 6. 功能架构表 (TBL1): remove 热点榜单 row, update 热点助手 row ----------
def iter_tables():
    return doc.tables

for tb in doc.tables:
    hdr = " ".join(c.text for c in tb.rows[0].cells)
    if hdr.startswith("一级模块"):
        # remove 热点榜单 row
        for ri in range(1, len(tb.rows)):
            rowtext = " ".join(c.text for c in tb.rows[ri].cells)
            if "热点榜单（已并入热点助手）" in rowtext:
                tbl_el = tb._tbl
                tr_el = tb.rows[ri]._tr
                tbl_el.remove(tr_el)
                break
        # update 热点助手 备注 cell
        for ri in range(1, len(tb.rows)):
            rowtext = " ".join(c.text for c in tb.rows[ri].cells)
            if "热点助手" in rowtext and "左上角 4 个页面导航" in rowtext:
                for c in tb.rows[ri].cells:
                    if "左上角 4 个页面导航" in c.text:
                        c.text = "左上角 4 个页面导航；其中国内热榜/国际热榜即原独立的「热点榜单」模块（国内16平台+国际5平台）"
                break
        break

# ---------- 7. 8.1 附录表 (TBL9): remove 热点榜单 row, relabel 热点助手 ----------
for tb in doc.tables:
    hdr = " ".join(c.text for c in tb.rows[0].cells)
    if hdr.startswith("页面") and "URL" in hdr:
        for ri in range(1, len(tb.rows)):
            rowtext = " ".join(c.text for c in tb.rows[ri].cells)
            if "热点榜单（已并入热点助手，/hotspot 重定向）" in rowtext:
                tbl_el = tb._tbl; tr_el = tb.rows[ri]._tr; tbl_el.remove(tr_el)
                break
        for ri in range(1, len(tb.rows)):
            rowtext = " ".join(c.text for c in tb.rows[ri].cells)
            if "热点助手（含左上角 4 个页面）" in rowtext:
                for c in tb.rows[ri].cells:
                    if "热点助手（含左上角 4 个页面）" in c.text:
                        c.text = "热点助手（4 个页面：热点搜索/热点预测/国内热榜/国际热榜）"
                break
        break

# ---------- 8. revision row V2.7 (TBL0) ----------
for tb in doc.tables:
    hdr = " ".join(c.text for c in tb.rows[0].cells)
    if hdr.startswith("版本") and "修订日期" in hdr:
        cells = tb.add_row().cells
        cells[0].text = "V2.7"
        cells[1].text = "2026-08-11"
        cells[2].text = ("按重新登录扫描的真实左侧导航架构重写第4章：移除独立的「热点榜单」模块章节，将其内容合并为热点助手的"
                         "「国内热榜/国际热榜」两个页面；热点助手调整为第4.3章并由左上角4个页面（热点搜索/热点预测/国内热榜/国际热榜）组成；"
                         "「历史记录」由独立章节归并至4.1新建任务的4.1.5；后续章节重新编号（生产工具→4.4、技能市场→4.5、个人中心→4.6）；"
                         "同步修订功能架构表、8.1附录、业务流程与非功能需求中的章节编号引用。")
        cells[3].text = "阿爪"
        break

doc.save(DST)
print("V2.7.docx saved ->", DST)
