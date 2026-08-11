from docx import Document
import glob, os

OLD="阿爪"; NEW="陈翃"

# ---------- 1) all PRD .md in docs/ ----------
print("=== MD ===")
for f in sorted(glob.glob("docs/数媒智创平台_PRD_V*.md")):
    with open(f, encoding="utf-8") as fh:
        t = fh.read()
    n = t.count(OLD)
    if n:
        t = t.replace(OLD, NEW)
        with open(f, "w", encoding="utf-8") as fh:
            fh.write(t)
    print(f"{os.path.basename(f)}: replaced {n}")

# ---------- 2) all PRD .docx (repo root) + roadmap docx ----------
print("=== DOCX ===")
def fix_docx(path):
    d = Document(path)
    cp = d.core_properties
    c = 0
    # body runs
    for p in d.paragraphs:
        for r in p.runs:
            if OLD in r.text:
                r.text = r.text.replace(OLD, NEW); c += r.text.count(NEW)
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if OLD in r.text:
                            r.text = r.text.replace(OLD, NEW); c += r.text.count(NEW)
                for t2 in cell.tables:
                    for r2 in t2.rows:
                        for cc in r2.cells:
                            for p in cc.paragraphs:
                                for r in p.runs:
                                    if OLD in r.text:
                                        r.text = r.text.replace(OLD, NEW); c += r.text.count(NEW)
    # author metadata
    meta_changed = False
    if (cp.author or "") != NEW:
        cp.author = NEW; meta_changed = True
    if (cp.last_modified_by or "") != NEW:
        cp.last_modified_by = NEW; meta_changed = True
    d.save(path)
    return c, meta_changed

for f in sorted(glob.glob("数媒智创平台_PRD_V*.docx")) + ["roadmap/数媒智创平台_2026H2产品规划.docx"]:
    c, mc = fix_docx(f)
    print(f"{f}: 阿爪 replaced {c}, author-meta set={mc}")
