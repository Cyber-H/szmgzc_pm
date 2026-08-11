# -*- coding: utf-8 -*-
"""修补 V2.5.docx -> V2.6.docx：热点助手导航形态由"顶部 Tab 栏/四个 Tab 整合页"
修正为"左上角 4 个页面"，并补充架构变迁说明；新增修订表 V2.6 行。"""
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell

SRC = r'C:\Users\chenhong\WorkBuddy\2026-06-26-14-49-33\szmgzc_pm\数媒智创平台_PRD_V2.5.docx'
DST = r'C:\Users\chenhong\WorkBuddy\2026-06-26-14-49-33\szmgzc_pm\数媒智创平台_PRD_V2.6.docx'

doc = Document(SRC)


def replace_in_runs(parent, old, new):
    paras = parent.paragraphs if hasattr(parent, 'paragraphs') else [parent]
    for p in paras:
        for r in p.runs:
            if old in r.text:
                r.text = r.text.replace(old, new)


def set_paragraph_text(p, new_text):
    runs = p.runs
    if not runs:
        p.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ''


def each_paragraph():
    yield from doc.paragraphs
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from cell.paragraphs


# A. 版本号
for p in doc.paragraphs:
    if '文档版本：V2.5' in p.text:
        replace_in_runs(p, '文档版本：V2.5', '文档版本：V2.6')

# B. 4.4.1 整段
NEW_441 = ('热点助手是基于热点数据的进阶热点分析工具，面向编辑、策划与运营人员。'
           'V2.4 起进行架构重组：原热点助手仅有「热点搜索 / 热点预测」2 个页面，'
           'V2.4 将原本独立的热点榜单模块并入，拆分为左上角并列的 4 个页面——'
           '热点搜索、热点预测、国内热榜、国际热榜（后两个即原热点榜单的国内/国际数据）。'
           '热点助手提供热点搜索、热点预测与事件深度分析能力，强调“查热点、跟热点、分析热点”，'
           '支持按行业、地域、时间等维度检索与预测热点事件。')
for p in doc.paragraphs:
    if p.text.startswith('热点助手是基于热点数据的进阶热点分析工具'):
        set_paragraph_text(p, NEW_441)
        break

# C. 4.4.2 第一行
for p in doc.paragraphs:
    if '热点助手 V2.4 升级为「四个 Tab 整合页」布局：' in p.text:
        replace_in_runs(p, '热点助手 V2.4 升级为「四个 Tab 整合页」布局：',
                        '热点助手 V2.4 重组后的页面结构如下：')

# D. 4.4.2 第二行
OLD_D = '顶部 Tab 栏：热点搜索 / 热点预测 / 国内热榜 / 国际热榜（后两个即原热点榜单的国内/国际数据）。'
NEW_D = ('左上角页面导航：4 个并列的页面入口——热点搜索 / 热点预测 / 国内热榜 / 国际热榜'
         '（后两个即原热点榜单的国内/国际数据，由原独立「热点榜单」模块并入）。')
for p in doc.paragraphs:
    if OLD_D in p.text:
        replace_in_runs(p, OLD_D, NEW_D)

# 通用 run-level 替换（覆盖段落 + 表格单元格文字）
REPLACES = [
    ('四个 Tab 整合页，原热点榜单现为其「国内热榜/国际热榜」Tab',
     '左上角 4 个页面导航，原热点榜单现为其「国内热榜/国际热榜」页面'),
    ('作为热点助手「国内热榜/国际热榜」Tab',
     '作为热点助手「国内热榜/国际热榜」页面'),
    ('热点助手（含四大 Tab）', '热点助手（含左上角 4 个页面）'),
    ('四个 Tab，原热点榜单的国内 16 平台',
     '左上角并列的 4 个页面，原热点榜单的国内 16 平台'),
    ('「国内热榜 / 国际热榜」Tab 呈现（平台源集合未变化）。',
     '「国内热榜 / 国际热榜」页面呈现（平台源集合未变化）。'),
    ('（热点搜索 / 热点预测 Tab 下）',
     '（热点搜索 / 热点预测 页面下）'),
    ('切换「国内热榜 / 国际热榜」Tab',
     '切换「国内热榜 / 国际热榜」页面'),
    ('国内/国际为两个独立 Tab',
     '国内/国际为两个独立页面'),
    ('均为「国内」Tab 平台；「国际」Tab 为独立国际热点源',
     '均为「国内」页面平台；「国际」页面为独立国际热点源'),
]
# 热点榜单平台源表的表头"Tab"列 -> "页面"
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            if cell.text.strip() == 'Tab':
                cell.text = '页面'
for p in each_paragraph():
    for old, new in REPLACES:
        if old in p.text:
            replace_in_runs(p, old, new)

# 修订表新增 V2.6 行
rev_table = None
for t in doc.tables:
    if 'V1.0' in t.rows[1].cells[0].text or '版本' in t.rows[0].cells[0].text:
        rev_table = t
        break

target_tr = None
for row in rev_table.rows:
    if row.cells[0].text.strip() == 'V2.5':
        target_tr = row._tr
        break

REV_CONTENT = ('热点助手架构描述修订（据用户反馈）：导航形态由“顶部 Tab 栏 / 四个 Tab 整合页”'
               '修正为“左上角 4 个页面”；并补充架构变迁说明——原热点助手仅有「热点搜索 / 热点预测」'
               '2 个页面，V2.4 起并入原独立「热点榜单」模块（国内热榜/国际热榜），拆分扩充为左上角并列的'
               '4 个页面。同步订正功能架构表、8.1 附录、第 9 章相关表述。')

if target_tr is not None:
    new_tr = copy.deepcopy(target_tr)
    target_tr.addnext(new_tr)
    new_tcs = new_tr.findall(qn('w:tc'))
    cells = [_Cell(tc, None) for tc in new_tcs]
    for c, txt in zip(cells, ['V2.6', '2026-08-11', REV_CONTENT, '阿爪']):
        c.text = txt
    print('revision row V2.6 added')
else:
    print('WARNING: V2.5 row not found in revision table')

doc.save(DST)
print('V2.6.docx generated ->', DST)
