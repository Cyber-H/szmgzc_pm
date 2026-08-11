# -*- coding: utf-8 -*-
"""Patch V2.4.docx -> V2.5.docx with the same textual edits applied to V2.5.md."""
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = r"C:\Users\chenhong\WorkBuddy\2026-06-26-14-49-33\szmgzc_pm\数媒智创平台_PRD_V2.4.docx"
DST = r"C:\Users\chenhong\WorkBuddy\2026-06-26-14-49-33\szmgzc_pm\数媒智创平台_PRD_V2.5.docx"

doc = Document(SRC)

def set_text(p_or_cell, new_text):
    if hasattr(p_or_cell, 'paragraphs'):
        p = p_or_cell.paragraphs[0]
    else:
        p = p_or_cell
    runs = p.runs
    if not runs:
        p.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""

def replace_in(doc, old, new):
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            set_text(p, p.text.replace(old, new))
            count += 1
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in para.text:
                        set_text(para, para.text.replace(old, new))
                        count += 1
    return count

def insert_paragraph_after(paragraph, text):
    """Insert a new paragraph (copying paragraph style) right after `paragraph`."""
    new_p = OxmlElement('w:p')
    pPr = paragraph._p.pPr
    if pPr is not None:
        new_p.append(copy.deepcopy(pPr))
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    p.add_run(text)
    return p

import re
def is_section_heading(p):
    s = p.text.strip()
    if re.match(r'^\d+(\.\d+){1,}\s', s):
        return True
    st = p.style
    return st is not None and st.name.lower().startswith('heading')

def rewrite_section_body(doc, heading_prefix, new_paragraphs):
    """Overwrite the body paragraphs of a numbered section (between its heading
    and the next heading) with `new_paragraphs`, preserving the heading itself.
    Robust to count mismatch: reuse existing real paragraphs, delete extras,
    insert if more are needed; stray blank paragraphs are removed."""
    paras = doc.paragraphs
    hi = next((i for i, p in enumerate(paras) if p.text.strip().startswith(heading_prefix)), None)
    if hi is None:
        print("!! heading not found:", heading_prefix)
        return
    ni = len(paras)
    for j in range(hi + 1, len(paras)):
        if is_section_heading(paras[j]):
            ni = j
            break
    body = paras[hi + 1:ni]
    real = [p for p in body if p.text.strip() != '']
    empties = [p for p in body if p.text.strip() == '']
    n, m = len(new_paragraphs), len(real)
    for k in range(min(n, m)):
        set_text(real[k], new_paragraphs[k])
    if n > m:
        anchor = real[m - 1] if m > 0 else paras[hi]
        for k in range(m, n):
            anchor = insert_paragraph_after(anchor, new_paragraphs[k])
    elif n < m:
        for k in range(n, m):
            p = real[k]
            p._p.getparent().remove(p._p)
    for p in empties:
        p._p.getparent().remove(p._p)

# ---- 1. version header ----
replace_in(doc, "文档版本：V2.4", "文档版本：V2.5")

# ---- 2. 功能架构表 热点助手核心功能 ----
replace_in(doc,
    "热点搜索、热点预测、国内/国际热榜、事件详情、热度趋势、事件脉络、深圳今日热点脉搏",
    "热点搜索、热点预测、国内/国际热榜、事件详情、热度趋势、事件脉络、深圳今日热点脉搏、一键生成热点专题（AI事件简报）")

# ---- 3. 4.4.5 事件详情：热度趋势段后插入「一键生成热点专题」 ----
NEW_P = ("一键生成热点专题：事件详情页底部新增「AI 事件简报（热点专题）」模块，提供「生成专题」按钮（V2.5 新增）。"
         "点击后一键基于当前事件异步生成一份 AI 事件简报（热点专题），生成任务在后台执行，不影响详情页其他内容浏览。"
         "每次生成的专题记录（含事件标题、生成时间、生成状态如“生成中”/“已完成”）汇总于模块下方的「任务历史」区，"
         "便于复盘与复用已生成的专题，支撑从“看热点”到“出专题报道”的内容生产闭环。")
for p in doc.paragraphs:
    if p.text.startswith("热度趋势：支持近24小时"):
        insert_paragraph_after(p, NEW_P)
        break

# ---- 4. 第9章：标题对齐 + 补全 9.2 / 9.3（V2.4 旧 docx 未写入 9.1/9.2，本次一并补齐） ----
replace_in(doc, "9. 本版线上核对说明（2026-07-14）", "9.1 2026-07-14 复核结论")

H92 = "9.2 2026-08-11 复核结论（V2.4）"
T92 = ("用户配合重新登录后再次逐模块扫描，重大结构更新如下：热点榜单 ↔ 热点助手 已合并（侧边栏不再有独立"
       "「热点榜单」菜单；访问 /hotspot 重定向至 /hotspot-assistant?tab=domestic；热点助手升级为"
       "「热点搜索 / 热点预测 / 国内热榜 / 国际热榜」四个 Tab，原热点榜单的国内 16 平台、国际 5 平台现作为"
       "「国内热榜 / 国际热榜」Tab 呈现，平台源集合未变化）；热点助手增强（检索区新增关键词+行业分类+地域+内容发布时间"
       "筛选与「刷新热点数据」「数据更新时间」，新增「深圳今日热点脉搏」地域化热点榜单带热度趋势↑，事件卡片升级展示"
       "分类/时间/地域/热度/阅读量并新增「所属中事件」关联）；知识库访问机制变更（主站不再以 iframe 内嵌 WeKnora，"
       "改由侧边栏菜单新标签跳转至 weknora.scms.sztv.com.cn 经 SSO 登录，主站 /weknora 直接访问返回「页面不存在」；"
       "WeKnora 内部优化升级内容因外部子站本环境无法访问，待用户真机补充）。"
       "待复核项（本环境仍无法访问外部子站/接口）：新建任务模型完整列表、生产工具（geekai）模型版本、"
       "技能市场（SkillHub）技能列表、热点助手事件详情页「事件脉络」数据接入情况。")

H93 = "9.3 2026-08-11 补充（V2.5）"
T93 = ("用户配合重新登录后再次扫描热点助手事件详情页，确认底部新增「AI 事件简报（热点专题）」模块与"
       "「生成专题」按钮（一键生成热点专题），生成记录统一管理于「任务历史」区。"
       "该能力完善了热点助手从热点发现到专题内容生产的闭环（详见 4.4.5 节）。")

for p in doc.paragraphs:
    if '热点助手事件详情页「事件脉络」数据是否已接入' in p.text:
        h92 = insert_paragraph_after(p, H92)
        c92 = insert_paragraph_after(h92, T92)
        h93 = insert_paragraph_after(c92, H93)
        insert_paragraph_after(h93, T93)
        break

# ---- 5. 修订记录表新增 V2.5 行 ----
def find_rev_table(doc):
    for t in doc.tables:
        txt = " ".join(c.text for c in t.rows[0].cells)
        if "版本" in txt and "修订日期" in txt:
            return t
    return None

rev = find_rev_table(doc)
if rev is not None:
    new_row = rev.add_row()
    vals = ["V2.5", "2026-08-11",
            "热点助手事件详情页新增「一键生成热点专题」能力：详情页底部新增「AI 事件简报（热点专题）」模块，"
            "提供「生成专题」按钮，点击后一键基于当前事件异步生成一份 AI 事件简报（热点专题），生成任务在后台执行；"
            "每次生成的专题记录（含事件标题、生成时间、生成状态如“生成中”/“已完成”）汇总于模块下方的「任务历史」区，"
            "便于复盘与复用已生成的专题。",
            "阿爪"]
    for i, cell in enumerate(new_row.cells):
        set_text(cell, vals[i])

# ---- 6. 4.4.2 页面结构：整节重写为与 md 一致（四个 Tab 整合页） ----
# 旧 docx 继承 V2.4 的过期结构（"顶部脉搏+中部筛选+下部列表"+"双Tab切换"），
# 当初替换字符串带多余 \n\n 未命中，此处按 md 重排整节。
SEC442 = [
    "热点助手 V2.4 升级为「四个 Tab 整合页」布局：",
    "顶部 Tab 栏：热点搜索 / 热点预测 / 国内热榜 / 国际热榜（后两个即原热点榜单的国内/国际数据）。",
    "检索区（热点搜索 / 热点预测 Tab 下）：关键词输入框 + 行业分类 + 地域 + 内容发布时间（起止）筛选 + 「搜索」「重置」按钮；V2.4 新增「刷新热点数据」按钮与「数据更新时间」标识，便于判断数据新鲜度。",
    "深圳今日热点脉搏：地域化热点榜单，展示深圳本地正在发酵的话题及热度上升趋势（如「深圳多点布置景观装置迎APEC ↑66」）。",
    "下部：热点事件卡片列表，每张卡片展示分类标签、地域、标题、摘要、热度、阅读量，并支持「所属中事件」关联跳转与「查看详情」。",
]
rewrite_section_body(doc, "4.4.2", SEC442)

doc.save(DST)
print("V2.5.docx generated ->", DST)
