# -*- coding: utf-8 -*-
"""Standalone fix: rewrite 4.4.2 body in V2.5.docx to match V2.5.md (four-Tab layout)."""
import copy, re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

PATH = r"C:\Users\chenhong\WorkBuddy\2026-06-26-14-49-33\szmgzc_pm\数媒智创平台_PRD_V2.5.docx"

doc = Document(PATH)

def set_text(p, new_text):
    runs = p.runs
    if not runs:
        p.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""

def insert_paragraph_after(paragraph, text):
    new_p = OxmlElement('w:p')
    pPr = paragraph._p.pPr
    if pPr is not None:
        new_p.append(copy.deepcopy(pPr))
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    p.add_run(text)
    return p

def is_section_heading(p):
    s = p.text.strip()
    if re.match(r'^\d+(\.\d+){1,}\s', s):
        return True
    st = p.style
    return st is not None and st.name.lower().startswith('heading')

def rewrite_section_body(doc, heading_prefix, new_paragraphs):
    paras = doc.paragraphs
    hi = next((i for i, p in enumerate(paras) if p.text.strip().startswith(heading_prefix)), None)
    if hi is None:
        print("!! heading not found:", heading_prefix)
        return
    ni = len(paras)
    for j in range(hi + 1, len(paras)):
        if is_section_heading(paras[j]):
            ni = j
            break
    body = paras[hi + 1:ni]
    real = [p for p in body if p.text.strip() != '']
    empties = [p for p in body if p.text.strip() == '']
    n, m = len(new_paragraphs), len(real)
    for k in range(min(n, m)):
        set_text(real[k], new_paragraphs[k])
    if n > m:
        anchor = real[m - 1] if m > 0 else paras[hi]
        for k in range(m, n):
            anchor = insert_paragraph_after(anchor, new_paragraphs[k])
    elif n < m:
        for k in range(n, m):
            p = real[k]
            p._p.getparent().remove(p._p)
    for p in empties:
        p._p.getparent().remove(p._p)

SEC442 = [
    "热点助手 V2.4 升级为「四个 Tab 整合页」布局：",
    "顶部 Tab 栏：热点搜索 / 热点预测 / 国内热榜 / 国际热榜（后两个即原热点榜单的国内/国际数据）。",
    "检索区（热点搜索 / 热点预测 Tab 下）：关键词输入框 + 行业分类 + 地域 + 内容发布时间（起止）筛选 + 「搜索」「重置」按钮；V2.4 新增「刷新热点数据」按钮与「数据更新时间」标识，便于判断数据新鲜度。",
    "深圳今日热点脉搏：地域化热点榜单，展示深圳本地正在发酵的话题及热度上升趋势（如「深圳多点布置景观装置迎APEC ↑66」）。",
    "下部：热点事件卡片列表，每张卡片展示分类标签、地域、标题、摘要、热度、阅读量，并支持「所属中事件」关联跳转与「查看详情」。",
]
rewrite_section_body(doc, "4.4.2", SEC442)

doc.save(PATH)
print("fixed 4.4.2 in", PATH)
