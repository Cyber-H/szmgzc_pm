import sys
from docx import Document
from docx.oxml.ns import qn

OLD = "阿爪"
NEW = "陈翃"

paths = [
    r"C:\Users\chenhong\WorkBuddy\2026-06-26-14-49-33\szmgzc_pm\数媒智创平台_PRD_V2.7.docx",
    r"C:\Users\chenhong\WorkBuddy\2026-06-26-14-49-33\数媒智创平台_PRD_V2.7.docx",
]

def replace_in_para(p):
    count = 0
    # Replace within individual runs (token is short / standalone -> single run)
    for run in p.runs:
        if OLD in run.text:
            run.text = run.text.replace(OLD, NEW)
            count += run.text.count(NEW)
    # Fallback: if still present at paragraph level (multi-run split), rebuild text
    if OLD in p.text:
        # collect runs, rebuild preserving first run formatting
        full = p.text
        full = full.replace(OLD, NEW)
        if p.runs:
            first = p.runs[0]
            first.text = full
            for r in p.runs[1:]:
                r.text = ""
            count += 1
    return count

total = 0
for path in paths:
    doc = Document(path)
    c = 0
    for p in doc.paragraphs:
        c += replace_in_para(p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    c += replace_in_para(p)
                # nested tables
                for t2 in cell.tables:
                    for r2 in t2.rows:
                        for cc in r2.cells:
                            for p in cc.paragraphs:
                                c += replace_in_para(p)
    doc.save(path)
    print(f"{path}: replaced {c} occurrence(s)")
    total += c

print("TOTAL replaced:", total)
